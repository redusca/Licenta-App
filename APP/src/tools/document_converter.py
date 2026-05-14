"""
Document Converter tool — convert between document formats locally.

Supported conversions
---------------------
PDF  → DOCX, TXT, HTML, PNG (renders each page as an image)
DOCX → PDF, TXT, HTML
HTML → PDF, DOCX
TXT  → PDF, DOCX
MD   → PDF, HTML, DOCX

Execution modes (outputMode)
-----------------------------
replace        : place the result alongside the original with the new extension.
copy           : place the result alongside the original with a _converted suffix.
virtual_drive  : copy results into a DocConvertResults virtual drive.
"""
from __future__ import annotations

import json
import os
import shutil
import subprocess
import uuid
import re
import tempfile
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path

from config import APP_VERSION
from migrations import get_latest_schema_version

# ── Paths ──────────────────────────────────────────────────────────────────────
_DATA_DIR = Path(__file__).parent.parent.parent / "data"
_TOOL_DRIVES_PATH = _DATA_DIR / "tool_drives.json"
_CONFIG_FILENAME = ".drive_config.json"

SUPPORTED_INPUT_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".html", ".htm", ".md", ".markdown"}

# Which formats each input type can convert to
CONVERSION_MAP: dict[str, list[str]] = {
    ".pdf":      ["docx", "txt", "html", "png"],
    ".docx":     ["pdf", "txt", "html"],
    ".doc":      ["pdf", "txt", "html"],
    ".txt":      ["pdf", "docx"],
    ".html":     ["pdf", "docx"],
    ".htm":      ["pdf", "docx"],
    ".md":       ["pdf", "html", "docx"],
    ".markdown": ["pdf", "html", "docx"],
}

# ── Agent tool definition ──────────────────────────────────────────────────────

DEFINITION = {
    "name": "document_converter",
    "description": (
        "Convert documents between formats: PDF, DOCX, TXT, HTML, Markdown, PNG. "
        "Supports batch processing and virtual-drive output."
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "files": {
                "type": "array",
                "description": 'List of file objects: [{"path": "...", "outputFormat": "pdf"}, ...]',
                "items": {"type": "object"},
            },
            "outputMode": {
                "type": "string",
                "enum": ["replace", "copy", "virtual_drive"],
                "description": "How to handle the output file(s).",
            },
            "outputPath": {
                "type": "string",
                "description": "Parent directory for virtual drive (only for virtual_drive mode).",
            },
        },
        "required": ["files", "outputMode"],
    },
    "input_instructions": (
        "files: array of {path, outputFormat} — use ask_user(input_type='file') to pick each document from a virtual drive. "
        "Supported inputs: PDF, DOCX, DOC, TXT, HTML, Markdown. "
        "outputFormat per file: 'pdf', 'docx', 'txt', 'html', or 'png' (PDF pages as images). "
        "outputMode: 'replace' places result alongside with new extension, 'copy' adds _converted suffix, 'virtual_drive' saves to a new virtual drive. "
        "outputPath: required only for virtual_drive — use ask_user(input_type='folder') to pick a folder from the app's virtual drives."
    ),
    "output_description": (
        "JSON {success, total, succeeded, failed, results:[{path, outputPath, success, pages?, error?}], virtualDrivePath?} "
        "— PDF→PNG produces a ZIP when the source has multiple pages."
    ),
}

# ── Tool-drives registry helpers ───────────────────────────────────────────────

def _load_tool_drives() -> list:
    if _TOOL_DRIVES_PATH.exists():
        try:
            return json.loads(_TOOL_DRIVES_PATH.read_text(encoding="utf-8"))
        except Exception:
            return []
    return []


def _save_tool_drives(drives: list) -> None:
    _DATA_DIR.mkdir(parents=True, exist_ok=True)
    _TOOL_DRIVES_PATH.write_text(json.dumps(drives, indent=2), encoding="utf-8")


def _register_tool_drive(drive_path: str, name: str, tool: str) -> None:
    drives = _load_tool_drives()
    normalized_new = os.path.normcase(os.path.normpath(drive_path))
    for d in drives:
        existing = d.get("path", "")
        if existing and os.path.normcase(os.path.normpath(existing)) == normalized_new:
            return
    drives.append({"path": drive_path, "name": name, "tool": tool})
    _save_tool_drives(drives)


def _ensure_virtual_drive(output_path: str) -> str:
    drive_name = "DocConvertResults"
    drive_path = os.path.join(output_path, drive_name)
    os.makedirs(drive_path, exist_ok=True)

    config_path = os.path.join(drive_path, _CONFIG_FILENAME)
    if not os.path.exists(config_path):
        config = {
            "schema_version": get_latest_schema_version(),
            "serial": str(uuid.uuid4()),
            "name": drive_name,
            "type": "move",
            "created_at": str(os.path.getctime(drive_path)),
            "app_version_created": APP_VERSION,
            "created_by_tool": "document_converter",
        }
        with open(config_path, "w", encoding="utf-8") as f:
            json.dump(config, f, indent=2)
        subprocess.call(["attrib", "+h", config_path], shell=True)

    _register_tool_drive(drive_path, drive_name, "document_converter")
    return drive_path


def _unique_path(path: str) -> str:
    """Return a non-colliding path by appending a numeric suffix if needed."""
    if not os.path.exists(path):
        return path
    stem, ext = os.path.splitext(path)
    counter = 1
    while os.path.exists(f"{stem}_{counter}{ext}"):
        counter += 1
    return f"{stem}_{counter}{ext}"


# ── Conversion functions ──────────────────────────────────────────────────────

def _pdf_to_docx(src: str, dst: str) -> str:
    from pdf2docx import Converter
    cv = Converter(src)
    cv.convert(dst)
    cv.close()
    return dst


def _docx_to_pdf(src: str, dst: str) -> str:
    from docx2pdf import convert
    convert(src, dst)
    return dst


def _pdf_to_txt(src: str, dst: str) -> str:
    from pypdf import PdfReader
    reader = PdfReader(src)
    text_parts = []
    for i, page in enumerate(reader.pages):
        page_text = page.extract_text() or ""
        text_parts.append(f"--- Page {i+1} ---\n{page_text}")
    with open(dst, "w", encoding="utf-8") as f:
        f.write("\n\n".join(text_parts))
    return dst


def _pdf_to_html(src: str, dst: str) -> str:
    """Convert PDF to HTML by extracting text and wrapping in styled HTML."""
    from pypdf import PdfReader
    reader = PdfReader(src)
    stem = Path(src).stem

    html_pages = []
    for i, page in enumerate(reader.pages):
        page_text = page.extract_text() or ""
        escaped = (
            page_text
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace("\n", "<br>\n")
        )
        html_pages.append(
            f'<div class="page" id="page-{i+1}">\n'
            f'  <h2>Page {i+1}</h2>\n'
            f'  <div class="content">{escaped}</div>\n'
            f'</div>'
        )

    html = (
        '<!DOCTYPE html>\n<html lang="en">\n<head>\n'
        f'  <meta charset="UTF-8">\n  <title>{stem}</title>\n'
        '  <style>\n'
        '    body { font-family: Georgia, serif; max-width: 800px; margin: 2rem auto; padding: 0 1rem; color: #333; }\n'
        '    .page { margin-bottom: 2rem; padding-bottom: 1rem; border-bottom: 1px solid #ddd; }\n'
        '    h2 { color: #555; font-size: 0.9rem; text-transform: uppercase; letter-spacing: 0.1em; }\n'
        '    .content { line-height: 1.6; white-space: pre-wrap; }\n'
        '  </style>\n'
        '</head>\n<body>\n'
        + "\n".join(html_pages) +
        '\n</body>\n</html>'
    )
    with open(dst, "w", encoding="utf-8") as f:
        f.write(html)
    return dst


def _pdf_to_png(src: str, dst_dir: str, stem: str) -> list[str]:
    """Render each PDF page as a PNG image. Returns list of output paths."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise RuntimeError(
            "PyMuPDF (fitz) is required for PDF→PNG conversion. "
            "Install it via: pip install PyMuPDF"
        )
    doc = fitz.open(src)
    outputs = []
    for i, page in enumerate(doc):
        pix = page.get_pixmap(dpi=200)
        out_path = _unique_path(os.path.join(dst_dir, f"{stem}_page_{i+1}.png"))
        pix.save(out_path)
        outputs.append(out_path)
    doc.close()
    return outputs


def _docx_to_txt(src: str, dst: str) -> str:
    """Extract plain text from DOCX."""
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError(
            "python-docx is required for DOCX→TXT conversion. "
            "Install it via: pip install python-docx"
        )
    doc = Document(src)
    text_parts = [p.text for p in doc.paragraphs]
    with open(dst, "w", encoding="utf-8") as f:
        f.write("\n".join(text_parts))
    return dst


def _docx_to_html(src: str, dst: str) -> str:
    """Convert DOCX to HTML using mammoth."""
    try:
        import mammoth
    except ImportError:
        raise RuntimeError(
            "mammoth is required for DOCX→HTML conversion. "
            "Install it via: pip install mammoth"
        )
    with open(src, "rb") as docx_file:
        result = mammoth.convert_to_html(docx_file)
    stem = Path(src).stem
    html = (
        '<!DOCTYPE html>\n<html lang="en">\n<head>\n'
        f'  <meta charset="UTF-8">\n  <title>{stem}</title>\n'
        '  <style>\n'
        '    body { font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif; '
        'max-width: 800px; margin: 2rem auto; padding: 0 1rem; color: #333; line-height: 1.6; }\n'
        '    img { max-width: 100%; }\n'
        '    table { border-collapse: collapse; width: 100%; }\n'
        '    td, th { border: 1px solid #ddd; padding: 8px; }\n'
        '  </style>\n'
        '</head>\n<body>\n'
        + result.value +
        '\n</body>\n</html>'
    )
    with open(dst, "w", encoding="utf-8") as f:
        f.write(html)
    return dst


def _txt_to_pdf(src: str, dst: str) -> str:
    """Convert plain text to PDF using reportlab."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm
        from reportlab.pdfgen import canvas
    except ImportError:
        raise RuntimeError(
            "reportlab is required for TXT→PDF conversion. "
            "Install it via: pip install reportlab"
        )
    with open(src, "r", encoding="utf-8", errors="replace") as f:
        lines = f.readlines()

    c = canvas.Canvas(dst, pagesize=A4)
    width, height = A4
    margin = 2 * cm
    y = height - margin
    line_height = 14
    c.setFont("Courier", 10)

    for line in lines:
        text = line.rstrip("\n\r")
        if y < margin:
            c.showPage()
            c.setFont("Courier", 10)
            y = height - margin
        c.drawString(margin, y, text)
        y -= line_height

    c.save()
    return dst


def _txt_to_docx(src: str, dst: str) -> str:
    """Convert plain text to DOCX."""
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError(
            "python-docx is required for TXT→DOCX conversion. "
            "Install it via: pip install python-docx"
        )
    with open(src, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()
    doc = Document()
    for para in text.split("\n"):
        doc.add_paragraph(para)
    doc.save(dst)
    return dst


def _html_to_pdf(src: str, dst: str) -> str:
    """Convert HTML to PDF using weasyprint or fallback to basic reportlab."""
    try:
        from weasyprint import HTML
        HTML(filename=src).write_pdf(dst)
        return dst
    except ImportError:
        pass

    # Fallback: strip tags and convert as text
    with open(src, "r", encoding="utf-8", errors="replace") as f:
        raw_html = f.read()
    text = re.sub(r'<[^>]+>', '', raw_html)
    text = text.strip()

    # Write text to temp file and use txt_to_pdf
    tmp_txt = dst + ".tmp.txt"
    try:
        with open(tmp_txt, "w", encoding="utf-8") as f:
            f.write(text)
        _txt_to_pdf(tmp_txt, dst)
    finally:
        if os.path.exists(tmp_txt):
            os.remove(tmp_txt)
    return dst


def _html_to_docx(src: str, dst: str) -> str:
    """Convert HTML to DOCX by extracting text and creating a Word document."""
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError(
            "python-docx is required for HTML→DOCX conversion. "
            "Install it via: pip install python-docx"
        )
    with open(src, "r", encoding="utf-8", errors="replace") as f:
        raw_html = f.read()

    # Basic HTML → text with paragraph splitting
    # Remove scripts and styles
    cleaned = re.sub(r'<(script|style)[^>]*>.*?</\1>', '', raw_html, flags=re.DOTALL | re.IGNORECASE)
    # Replace <br>, <p>, </p>, <div>, etc. with newlines
    cleaned = re.sub(r'<br\s*/?\s*>', '\n', cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r'</?(p|div|h[1-6]|li|tr)[^>]*>', '\n', cleaned, flags=re.IGNORECASE)
    # Strip remaining tags
    text = re.sub(r'<[^>]+>', '', cleaned)
    # Clean up excess whitespace
    text = re.sub(r'\n{3,}', '\n\n', text).strip()

    doc = Document()
    for para in text.split("\n"):
        stripped = para.strip()
        if stripped:
            doc.add_paragraph(stripped)
    doc.save(dst)
    return dst


def _md_to_html(src: str, dst: str) -> str:
    """Convert Markdown to HTML."""
    try:
        import markdown
    except ImportError:
        raise RuntimeError(
            "markdown is required for MD→HTML conversion. "
            "Install it via: pip install markdown"
        )
    with open(src, "r", encoding="utf-8", errors="replace") as f:
        md_text = f.read()

    html_body = markdown.markdown(md_text, extensions=["tables", "fenced_code", "toc"])
    stem = Path(src).stem
    html = (
        '<!DOCTYPE html>\n<html lang="en">\n<head>\n'
        f'  <meta charset="UTF-8">\n  <title>{stem}</title>\n'
        '  <style>\n'
        '    body { font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif; '
        'max-width: 800px; margin: 2rem auto; padding: 0 1rem; color: #333; line-height: 1.7; }\n'
        '    code { background: #f4f4f4; padding: 2px 6px; border-radius: 3px; font-size: 0.9em; }\n'
        '    pre { background: #f4f4f4; padding: 1rem; border-radius: 6px; overflow-x: auto; }\n'
        '    pre code { background: none; padding: 0; }\n'
        '    table { border-collapse: collapse; width: 100%; margin: 1rem 0; }\n'
        '    td, th { border: 1px solid #ddd; padding: 8px 12px; }\n'
        '    th { background: #f8f8f8; }\n'
        '    blockquote { border-left: 3px solid #ddd; margin-left: 0; padding-left: 1rem; color: #666; }\n'
        '    img { max-width: 100%; }\n'
        '  </style>\n'
        '</head>\n<body>\n'
        + html_body +
        '\n</body>\n</html>'
    )
    with open(dst, "w", encoding="utf-8") as f:
        f.write(html)
    return dst


def _md_to_pdf(src: str, dst: str) -> str:
    """Convert Markdown → HTML → PDF."""
    tmp_html = dst + ".tmp.html"
    try:
        _md_to_html(src, tmp_html)
        _html_to_pdf(tmp_html, dst)
    finally:
        if os.path.exists(tmp_html):
            os.remove(tmp_html)
    return dst


def _md_to_docx(src: str, dst: str) -> str:
    """Convert Markdown → HTML → DOCX."""
    tmp_html = dst + ".tmp.html"
    try:
        _md_to_html(src, tmp_html)
        _html_to_docx(tmp_html, dst)
    finally:
        if os.path.exists(tmp_html):
            os.remove(tmp_html)
    return dst


# ── Conversion router ─────────────────────────────────────────────────────────

def _get_converter(src_ext: str, target_fmt: str):
    """Return (converter_fn, output_ext) or raise ValueError."""
    src_ext = src_ext.lower()
    target_fmt = target_fmt.lower()

    _CONVERTERS = {
        (".pdf", "docx"):     (_pdf_to_docx, ".docx"),
        (".pdf", "txt"):      (_pdf_to_txt, ".txt"),
        (".pdf", "html"):     (_pdf_to_html, ".html"),
        (".pdf", "png"):      (None, ".png"),  # special multi-output
        (".docx", "pdf"):     (_docx_to_pdf, ".pdf"),
        (".docx", "txt"):     (_docx_to_txt, ".txt"),
        (".docx", "html"):    (_docx_to_html, ".html"),
        (".doc", "pdf"):      (_docx_to_pdf, ".pdf"),
        (".doc", "txt"):      (_docx_to_txt, ".txt"),
        (".doc", "html"):     (_docx_to_html, ".html"),
        (".txt", "pdf"):      (_txt_to_pdf, ".pdf"),
        (".txt", "docx"):     (_txt_to_docx, ".docx"),
        (".html", "pdf"):     (_html_to_pdf, ".pdf"),
        (".html", "docx"):    (_html_to_docx, ".docx"),
        (".htm", "pdf"):      (_html_to_pdf, ".pdf"),
        (".htm", "docx"):     (_html_to_docx, ".docx"),
        (".md", "pdf"):       (_md_to_pdf, ".pdf"),
        (".md", "html"):      (_md_to_html, ".html"),
        (".md", "docx"):      (_md_to_docx, ".docx"),
        (".markdown", "pdf"): (_md_to_pdf, ".pdf"),
        (".markdown", "html"):(_md_to_html, ".html"),
        (".markdown", "docx"):(_md_to_docx, ".docx"),
    }

    key = (src_ext, target_fmt)
    if key not in _CONVERTERS:
        raise ValueError(f"Unsupported conversion: {src_ext} → .{target_fmt}")
    return _CONVERTERS[key]


# ── Single-item processor ─────────────────────────────────────────────────────

def _process_single_item(
    item: dict,
    output_mode: str,
    virtual_drive_path: str | None,
) -> dict:
    src = item.get("path", "")
    target_fmt = item.get("outputFormat", "").lower().lstrip(".")

    if not src or not os.path.isfile(src):
        return {"path": src, "success": False, "error": "File not found"}

    src_ext = Path(src).suffix.lower()
    stem = Path(src).stem

    if src_ext not in SUPPORTED_INPUT_EXTENSIONS:
        return {"path": src, "success": False, "error": f"Unsupported input format: {src_ext}"}

    allowed = CONVERSION_MAP.get(src_ext, [])
    if target_fmt not in allowed:
        return {"path": src, "success": False, "error": f"Cannot convert {src_ext} → .{target_fmt}"}

    try:
        converter_fn, out_ext = _get_converter(src_ext, target_fmt)

        # Special case: PDF → PNG (multi-output → zip if multiple pages)
        if src_ext == ".pdf" and target_fmt == "png":
            if output_mode == "virtual_drive" and virtual_drive_path:
                out_dir = virtual_drive_path
            elif output_mode == "copy":
                out_dir = os.path.dirname(src)
            else:  # replace
                out_dir = os.path.dirname(src)

            outputs = _pdf_to_png(src, out_dir, stem)

            # Single page → just return the single PNG
            if len(outputs) <= 1:
                return {
                    "path": src,
                    "outputPath": outputs[0] if outputs else "",
                    "success": True,
                    "pages": len(outputs),
                }

            # Multiple pages → bundle into a ZIP and remove individual PNGs
            import zipfile
            zip_path = _unique_path(os.path.join(out_dir, f"{stem}_pages.zip"))
            with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
                for png_path in outputs:
                    zf.write(png_path, os.path.basename(png_path))
            # Clean up individual PNGs
            for png_path in outputs:
                try:
                    os.remove(png_path)
                except OSError:
                    pass

            return {
                "path": src,
                "outputPath": zip_path,
                "success": True,
                "pages": len(outputs),
            }

        # Determine output path
        if output_mode == "replace":
            final = os.path.join(os.path.dirname(src), f"{stem}{out_ext}")
            final = _unique_path(final)
            converter_fn(src, final)

        elif output_mode == "copy":
            candidate = os.path.join(os.path.dirname(src), f"{stem}_converted{out_ext}")
            final = _unique_path(candidate)
            converter_fn(src, final)

        else:  # virtual_drive
            dest = os.path.join(virtual_drive_path, f"{stem}{out_ext}")  # type: ignore[arg-type]
            dest = _unique_path(dest)
            converter_fn(src, dest)
            final = dest

        return {"path": src, "outputPath": final, "success": True}

    except Exception as exc:
        return {"path": src, "success": False, "error": str(exc)}


# ── Public API ─────────────────────────────────────────────────────────────────

def execute(input_data: dict) -> str:
    """Synchronous single-threaded execution."""
    return execute_parallel(input_data, max_workers=1)


def execute_parallel(input_data: dict, max_workers: int = 2) -> str:
    files: list = input_data.get("files", [])
    output_mode: str = input_data.get("outputMode", "copy")
    output_path: str = input_data.get("outputPath", "")

    if not files:
        return json.dumps({"success": False, "error": "No files provided.", "results": []})

    virtual_drive_path: str | None = None
    if output_mode == "virtual_drive":
        if not output_path or not os.path.isdir(output_path):
            return json.dumps({
                "success": False,
                "error": f"Invalid output path for virtual drive: '{output_path}'",
                "results": [],
            })
        virtual_drive_path = _ensure_virtual_drive(output_path)

    workers = 1 if output_mode == "replace" else min(max_workers, len(files))

    results: list = [None] * len(files)
    with ThreadPoolExecutor(max_workers=workers) as pool:
        future_to_idx = {
            pool.submit(
                _process_single_item, item, output_mode, virtual_drive_path,
            ): idx
            for idx, item in enumerate(files)
        }
        for future in as_completed(future_to_idx):
            idx = future_to_idx[future]
            try:
                results[idx] = future.result()
            except Exception as exc:
                results[idx] = {"path": files[idx].get("path", ""), "success": False, "error": str(exc)}

    succeeded = sum(1 for r in results if r and r["success"])
    response: dict = {
        "success": succeeded > 0,
        "total": len(results),
        "succeeded": succeeded,
        "failed": len(results) - succeeded,
        "results": results,
    }
    if virtual_drive_path:
        response["virtualDrivePath"] = virtual_drive_path

    return json.dumps(response)
